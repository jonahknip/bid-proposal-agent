"""
Report Generator - Generate Word documents and Excel spreadsheets for bid analysis
"""

import io
import os
from typing import Dict, Any, List, Optional
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


class ReportGenerator:
    """
    Generates bid analysis reports in Word and Excel formats.
    Matches Abonmarche styling and formatting.
    """
    
    # Brand colors
    NAVY = RGBColor(27, 54, 93)  # #1B365D
    RED = RGBColor(200, 16, 46)  # #C8102E
    GREEN = RGBColor(40, 167, 69)
    ORANGE = RGBColor(230, 126, 34)
    GRAY = RGBColor(108, 117, 125)
    
    # Excel colors (hex)
    EXCEL_NAVY = "1B365D"
    EXCEL_RED = "C8102E"
    EXCEL_GREEN = "28A745"
    EXCEL_ORANGE = "E67E22"
    EXCEL_GRAY = "6C757D"
    EXCEL_LIGHT_GRAY = "F5F5F5"
    
    def __init__(self):
        """Initialize the report generator."""
        pass
    
    def _set_cell_shading(self, cell, color_hex: str):
        """Set background shading for a table cell."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def generate_bid_analysis_report(self, analysis: Dict[str, Any], project_name: str = "") -> io.BytesIO:
        """
        Generate a comprehensive Word document report for bid analysis.
        
        Args:
            analysis: Complete analysis results from BidAnalyzer
            project_name: Project name for the report
            
        Returns:
            BytesIO buffer containing the Word document
        """
        doc = Document()
        
        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        
        # Title
        title = doc.add_heading('BID PROPOSAL ANALYSIS REPORT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = self.NAVY
        
        # Subtitle with project name
        if project_name:
            subtitle = doc.add_paragraph(project_name)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle.runs[0]
            run.font.size = Pt(14)
            run.font.color.rgb = self.RED
        
        # Date
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Executive Summary
        self._add_section_header(doc, "EXECUTIVE SUMMARY")
        
        summary = analysis.get('summary', {})
        status = self._get_status_info(analysis)
        
        # Status box
        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f"STATUS: {status['status']}")
        status_run.bold = True
        status_run.font.size = Pt(14)
        if status['color'] == 'green':
            status_run.font.color.rgb = self.GREEN
        elif status['color'] == 'red':
            status_run.font.color.rgb = self.RED
        else:
            status_run.font.color.rgb = self.ORANGE
        
        status_para.add_run(f"\n{status['message']}")
        
        # Scores table
        scores_table = doc.add_table(rows=1, cols=4)
        scores_table.style = 'Table Grid'
        
        headers = ['Completeness', 'Accuracy', 'Critical Issues', 'Warnings']
        values = [
            f"{status.get('completeness_score', 0):.0f}%",
            f"{status.get('accuracy_score', 0):.0f}%",
            str(status.get('critical_issues', 0)),
            str(status.get('warnings', 0))
        ]
        
        row = scores_table.rows[0]
        for i, (header, value) in enumerate(zip(headers, values)):
            cell = row.cells[i]
            cell.text = f"{header}\n{value}"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Critical Issues
        critical = analysis.get('critical_issues', [])
        if critical:
            self._add_section_header(doc, "CRITICAL ISSUES", color=self.RED)
            for issue in critical:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(issue)
                run.font.color.rgb = self.RED
            doc.add_paragraph()
        
        # Line Item Analysis
        self._add_section_header(doc, "LINE ITEM ANALYSIS")
        
        rule_based = analysis.get('rule_based', {})
        
        # Matches
        matches = rule_based.get('matches', [])
        if matches:
            doc.add_paragraph().add_run('Matching Items').bold = True
            self._add_quantity_table(doc, matches, 'match')
        
        # Discrepancies
        discrepancies = rule_based.get('discrepancies', [])
        if discrepancies:
            doc.add_paragraph().add_run('Quantity Discrepancies').bold = True
            self._add_quantity_table(doc, discrepancies, 'discrepancy')
        
        # Missing Items
        missing = rule_based.get('missing', [])
        if missing:
            doc.add_paragraph().add_run('Missing Items (Required but not in proposal)').bold = True
            self._add_missing_items_table(doc, missing)
        
        # Extra Items
        extra = rule_based.get('extra', [])
        if extra:
            doc.add_paragraph().add_run('Extra Items (In proposal but not required)').bold = True
            self._add_missing_items_table(doc, extra)
        
        # Recommendations
        recommendations = analysis.get('recommendations', [])
        if recommendations:
            self._add_section_header(doc, "RECOMMENDATIONS")
            for i, rec in enumerate(recommendations, 1):
                para = doc.add_paragraph()
                para.add_run(f"{i}. ").bold = True
                para.add_run(rec)
        
        # Warnings
        warnings = analysis.get('warnings', [])
        if warnings:
            self._add_section_header(doc, "WARNINGS")
            for warning in warnings:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(warning)
                run.font.color.rgb = self.ORANGE
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Generated by Bid Proposal Agent - Abonmarche")
        run.font.size = Pt(9)
        run.font.color.rgb = self.GRAY
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _add_section_header(self, doc, text: str, color=None):
        """Add a styled section header."""
        if color is None:
            color = self.NAVY
        
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.color.rgb = color
        
        # Add underline
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
    
    def _add_quantity_table(self, doc, items: List[Dict], item_type: str = 'match'):
        """Add a table for quantity items."""
        if not items:
            return
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Description', 'Required', 'Proposed', 'Unit', 'Variance']
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            self._set_cell_shading(cell, self.EXCEL_NAVY)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        for item in items:
            row = table.add_row()
            row.cells[0].text = str(item.get('description', ''))[:50]
            row.cells[1].text = str(item.get('required_qty', item.get('quantity', '')))
            row.cells[2].text = str(item.get('proposed_qty', ''))
            row.cells[3].text = str(item.get('unit', ''))
            
            variance = item.get('variance_pct', 0)
            row.cells[4].text = f"{variance:+.1f}%" if variance else "0%"
            
            # Color code variance
            if item_type == 'discrepancy' or abs(variance) > 5:
                if variance > 0:
                    row.cells[4].paragraphs[0].runs[0].font.color.rgb = self.ORANGE
                else:
                    row.cells[4].paragraphs[0].runs[0].font.color.rgb = self.RED
        
        doc.add_paragraph()
    
    def _add_missing_items_table(self, doc, items: List[Dict]):
        """Add a table for missing/extra items."""
        if not items:
            return
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['Item No.', 'Description', 'Quantity', 'Unit']
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            self._set_cell_shading(cell, self.EXCEL_NAVY)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        for item in items:
            row = table.add_row()
            row.cells[0].text = str(item.get('item_number', ''))
            row.cells[1].text = str(item.get('description', ''))[:50]
            row.cells[2].text = str(item.get('quantity', ''))
            row.cells[3].text = str(item.get('unit', ''))
        
        doc.add_paragraph()
    
    def _get_status_info(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Get status information from analysis."""
        summary = analysis.get('summary', {})
        rule_based = analysis.get('rule_based', {})
        
        return {
            'status': summary.get('recommendation', 'REVIEW'),
            'color': 'green' if summary.get('recommendation') == 'go' else 'orange',
            'message': summary.get('overall_assessment', ''),
            'completeness_score': summary.get('completeness_score', 0) or rule_based.get('completeness_score', 0),
            'accuracy_score': summary.get('accuracy_score', 0) or rule_based.get('accuracy_score', 0),
            'critical_issues': len(analysis.get('critical_issues', [])),
            'warnings': len(analysis.get('warnings', []))
        }
    
    def generate_quantity_excel(
        self,
        quantities: List[Dict[str, Any]],
        project_name: str = "",
        include_comparison: bool = False,
        comparison_data: Optional[Dict[str, Any]] = None
    ) -> io.BytesIO:
        """
        Generate an Excel spreadsheet with quantities.
        
        Args:
            quantities: List of quantity items
            project_name: Project name for the header
            include_comparison: Whether to include comparison columns
            comparison_data: Data for quantity comparison
            
        Returns:
            BytesIO buffer containing the Excel file
        """
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Quantities"
        
        # Styling
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color=self.EXCEL_NAVY, end_color=self.EXCEL_NAVY, fill_type="solid")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Title
        ws.merge_cells('A1:G1')
        ws['A1'] = f"QUANTITY SUMMARY - {project_name}" if project_name else "QUANTITY SUMMARY"
        ws['A1'].font = Font(bold=True, size=14, color=self.EXCEL_NAVY)
        ws['A1'].alignment = Alignment(horizontal='center')
        
        ws['A2'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        ws['A2'].font = Font(size=9, color=self.EXCEL_GRAY)
        
        # Headers
        headers = ['Item No.', 'Description', 'Quantity', 'Unit', 'Category', 'Subcategory', 'Sheet Ref.']
        if include_comparison:
            headers.extend(['Plan Qty', 'Variance', 'Variance %'])
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', wrap_text=True)
        
        # Data rows
        row_num = 5
        for qty in quantities:
            ws.cell(row=row_num, column=1, value=qty.get('item_number', '')).border = thin_border
            ws.cell(row=row_num, column=2, value=qty.get('description', '')).border = thin_border
            
            qty_cell = ws.cell(row=row_num, column=3, value=qty.get('quantity', 0))
            qty_cell.border = thin_border
            qty_cell.number_format = '#,##0.00'
            
            ws.cell(row=row_num, column=4, value=qty.get('unit', '')).border = thin_border
            ws.cell(row=row_num, column=5, value=qty.get('category', '')).border = thin_border
            ws.cell(row=row_num, column=6, value=qty.get('subcategory', '')).border = thin_border
            ws.cell(row=row_num, column=7, value=qty.get('sheet_reference', '')).border = thin_border
            
            if include_comparison and comparison_data:
                plan_qty = comparison_data.get(qty.get('description', ''), {}).get('plan_qty', 0)
                variance = qty.get('quantity', 0) - plan_qty
                variance_pct = (variance / plan_qty * 100) if plan_qty else 0
                
                ws.cell(row=row_num, column=8, value=plan_qty).border = thin_border
                
                var_cell = ws.cell(row=row_num, column=9, value=variance)
                var_cell.border = thin_border
                var_cell.number_format = '+#,##0.00;-#,##0.00;0'
                if variance < 0:
                    var_cell.font = Font(color=self.EXCEL_RED)
                elif variance > 0:
                    var_cell.font = Font(color=self.EXCEL_ORANGE)
                
                var_pct_cell = ws.cell(row=row_num, column=10, value=variance_pct / 100)
                var_pct_cell.border = thin_border
                var_pct_cell.number_format = '+0.0%;-0.0%;0%'
            
            row_num += 1
        
        # Adjust column widths
        column_widths = [12, 45, 12, 8, 15, 15, 25]
        if include_comparison:
            column_widths.extend([12, 12, 12])
        
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width
        
        # Category Summary Sheet
        ws_summary = wb.create_sheet("Summary by Category")
        self._add_category_summary_sheet(ws_summary, quantities)
        
        # Save to buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _add_category_summary_sheet(self, ws, quantities: List[Dict[str, Any]]):
        """Add a summary sheet organized by category."""
        
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color=self.EXCEL_NAVY, end_color=self.EXCEL_NAVY, fill_type="solid")
        category_fill = PatternFill(start_color=self.EXCEL_LIGHT_GRAY, end_color=self.EXCEL_LIGHT_GRAY, fill_type="solid")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Title
        ws['A1'] = "QUANTITY SUMMARY BY CATEGORY"
        ws['A1'].font = Font(bold=True, size=14, color=self.EXCEL_NAVY)
        
        # Organize by category
        categories = {}
        for qty in quantities:
            cat = qty.get('category', 'Misc')
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(qty)
        
        row_num = 3
        
        for category, items in sorted(categories.items()):
            # Category header
            ws.cell(row=row_num, column=1, value=category.upper().replace('_', ' ')).font = Font(bold=True, size=12)
            ws.cell(row=row_num, column=1).fill = category_fill
            ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=4)
            row_num += 1
            
            # Column headers
            for col, header in enumerate(['Description', 'Quantity', 'Unit', 'Subcategory'], 1):
                cell = ws.cell(row=row_num, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
            row_num += 1
            
            # Items
            for item in items:
                ws.cell(row=row_num, column=1, value=item.get('description', '')).border = thin_border
                qty_cell = ws.cell(row=row_num, column=2, value=item.get('quantity', 0))
                qty_cell.border = thin_border
                qty_cell.number_format = '#,##0.00'
                ws.cell(row=row_num, column=3, value=item.get('unit', '')).border = thin_border
                ws.cell(row=row_num, column=4, value=item.get('subcategory', '')).border = thin_border
                row_num += 1
            
            row_num += 1  # Space between categories
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 45
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 20
    
    def generate_comparison_excel(
        self,
        comparison_results: Dict[str, Any],
        project_name: str = ""
    ) -> io.BytesIO:
        """
        Generate an Excel report comparing proposal quantities to plan quantities.
        
        Args:
            comparison_results: Results from BidAnalyzer.compare_quantities()
            project_name: Project name for the header
            
        Returns:
            BytesIO buffer containing the Excel file
        """
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Comparison"
        
        # Styling
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color=self.EXCEL_NAVY, end_color=self.EXCEL_NAVY, fill_type="solid")
        green_fill = PatternFill(start_color="D4EDDA", end_color="D4EDDA", fill_type="solid")
        red_fill = PatternFill(start_color="F8D7DA", end_color="F8D7DA", fill_type="solid")
        orange_fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Title
        ws.merge_cells('A1:G1')
        ws['A1'] = f"QUANTITY COMPARISON - {project_name}" if project_name else "QUANTITY COMPARISON"
        ws['A1'].font = Font(bold=True, size=14, color=self.EXCEL_NAVY)
        
        # Summary
        summary = comparison_results.get('summary', {})
        ws['A3'] = f"Match Rate: {summary.get('match_rate', 0):.1f}%"
        ws['A3'].font = Font(bold=True)
        ws['C3'] = f"Matches: {summary.get('matches', 0)}"
        ws['E3'] = f"Over: {summary.get('over_estimated', 0)}"
        ws['E3'].font = Font(color=self.EXCEL_ORANGE)
        ws['G3'] = f"Under: {summary.get('under_estimated', 0)}"
        ws['G3'].font = Font(color=self.EXCEL_RED)
        
        # Headers
        headers = ['Description', 'Proposal Qty', 'Plan Qty', 'Unit', 'Difference', 'Variance %', 'Status']
        row_num = 5
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row_num, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center')
        
        row_num += 1
        
        # Matches
        for item in comparison_results.get('matches', []):
            self._add_comparison_row(ws, row_num, item, 'MATCH', green_fill, thin_border)
            row_num += 1
        
        # Over-estimated
        for item in comparison_results.get('over_estimated', []):
            self._add_comparison_row(ws, row_num, item, 'OVER', orange_fill, thin_border)
            row_num += 1
        
        # Under-estimated
        for item in comparison_results.get('under_estimated', []):
            self._add_comparison_row(ws, row_num, item, 'UNDER', red_fill, thin_border)
            row_num += 1
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 40
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 10
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 12
        ws.column_dimensions['G'].width = 10
        
        # Missing Items Sheet
        if comparison_results.get('not_in_proposal'):
            ws_missing = wb.create_sheet("Missing in Proposal")
            self._add_simple_items_sheet(
                ws_missing, 
                comparison_results['not_in_proposal'],
                "Items on Plans but Missing from Proposal"
            )
        
        # Extra Items Sheet  
        if comparison_results.get('not_on_plans'):
            ws_extra = wb.create_sheet("Not on Plans")
            self._add_simple_items_sheet(
                ws_extra,
                comparison_results['not_on_plans'],
                "Items in Proposal but Not Found on Plans"
            )
        
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _add_comparison_row(self, ws, row: int, item: Dict, status: str, fill, border):
        """Add a row to the comparison sheet."""
        ws.cell(row=row, column=1, value=item.get('description', '')).border = border
        
        prop_cell = ws.cell(row=row, column=2, value=item.get('proposal_qty', 0))
        prop_cell.border = border
        prop_cell.number_format = '#,##0.00'
        
        plan_cell = ws.cell(row=row, column=3, value=item.get('plan_qty', 0))
        plan_cell.border = border
        plan_cell.number_format = '#,##0.00'
        
        ws.cell(row=row, column=4, value=item.get('unit', '')).border = border
        
        diff_cell = ws.cell(row=row, column=5, value=item.get('difference', 0))
        diff_cell.border = border
        diff_cell.number_format = '+#,##0.00;-#,##0.00;0'
        
        var_cell = ws.cell(row=row, column=6, value=item.get('variance_pct', 0) / 100)
        var_cell.border = border
        var_cell.number_format = '+0.0%;-0.0%;0%'
        
        status_cell = ws.cell(row=row, column=7, value=status)
        status_cell.border = border
        status_cell.fill = fill
        status_cell.alignment = Alignment(horizontal='center')
    
    def _add_simple_items_sheet(self, ws, items: List[Dict], title: str):
        """Add a simple items list sheet."""
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color=self.EXCEL_NAVY, end_color=self.EXCEL_NAVY, fill_type="solid")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        ws['A1'] = title
        ws['A1'].font = Font(bold=True, size=12, color=self.EXCEL_NAVY)
        
        headers = ['Description', 'Quantity', 'Unit', 'Category']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
        
        row_num = 4
        for item in items:
            ws.cell(row=row_num, column=1, value=item.get('description', '')).border = thin_border
            ws.cell(row=row_num, column=2, value=item.get('quantity', 0)).border = thin_border
            ws.cell(row=row_num, column=3, value=item.get('unit', '')).border = thin_border
            ws.cell(row=row_num, column=4, value=item.get('category', '')).border = thin_border
            row_num += 1
        
        ws.column_dimensions['A'].width = 45
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 15
    
    def generate_html_report(self, analysis: Dict[str, Any], project_name: str = "") -> str:
        """
        Generate an HTML report for display in the web UI.
        
        Args:
            analysis: Complete analysis results
            project_name: Project name
            
        Returns:
            HTML string
        """
        status = self._get_status_info(analysis)
        summary = analysis.get('summary', {})
        rule_based = analysis.get('rule_based', {})
        
        status_color = '#28a745' if status['color'] == 'green' else '#dc3545' if status['color'] == 'red' else '#ffc107'
        
        html = f'''
<div class="bid-analysis-report">
    <div class="report-header" style="background: #1B365D; color: white; padding: 20px; border-radius: 8px 8px 0 0;">
        <h1 style="margin: 0;">Bid Proposal Analysis</h1>
        <p style="margin: 5px 0 0 0; opacity: 0.9;">{project_name or 'Project Analysis'}</p>
    </div>
    
    <div class="status-banner" style="background: {status_color}20; border-left: 4px solid {status_color}; padding: 15px; margin: 20px 0;">
        <h2 style="color: {status_color}; margin: 0;">Status: {status['status']}</h2>
        <p style="margin: 10px 0 0 0;">{status['message']}</p>
    </div>
    
    <div class="scores-grid" style="display: grid; grid-template-columns: repeat(4, 1fr); gap: 15px; margin: 20px 0;">
        <div class="score-box" style="background: #e3f2fd; padding: 15px; border-radius: 8px; text-align: center;">
            <div style="font-size: 24px; font-weight: bold; color: #1B365D;">{status.get('completeness_score', 0):.0f}%</div>
            <div style="font-size: 12px; color: #666;">Completeness</div>
        </div>
        <div class="score-box" style="background: #e8f5e9; padding: 15px; border-radius: 8px; text-align: center;">
            <div style="font-size: 24px; font-weight: bold; color: #28a745;">{status.get('accuracy_score', 0):.0f}%</div>
            <div style="font-size: 12px; color: #666;">Accuracy</div>
        </div>
        <div class="score-box" style="background: #ffebee; padding: 15px; border-radius: 8px; text-align: center;">
            <div style="font-size: 24px; font-weight: bold; color: #dc3545;">{status.get('critical_issues', 0)}</div>
            <div style="font-size: 12px; color: #666;">Critical Issues</div>
        </div>
        <div class="score-box" style="background: #fff3e0; padding: 15px; border-radius: 8px; text-align: center;">
            <div style="font-size: 24px; font-weight: bold; color: #ff9800;">{status.get('warnings', 0)}</div>
            <div style="font-size: 12px; color: #666;">Warnings</div>
        </div>
    </div>
'''
        
        # Critical Issues
        critical = analysis.get('critical_issues', [])
        if critical:
            html += '''
    <div class="section" style="margin: 20px 0;">
        <h3 style="color: #dc3545; border-bottom: 2px solid #dc3545; padding-bottom: 8px;">Critical Issues</h3>
        <ul style="list-style: none; padding: 0;">
'''
            for issue in critical:
                html += f'            <li style="padding: 10px; background: #ffebee; margin: 5px 0; border-radius: 4px; border-left: 4px solid #dc3545;">{issue}</li>\n'
            html += '''        </ul>
    </div>
'''
        
        # Discrepancies
        discrepancies = rule_based.get('discrepancies', [])
        if discrepancies:
            html += '''
    <div class="section" style="margin: 20px 0;">
        <h3 style="color: #1B365D; border-bottom: 2px solid #1B365D; padding-bottom: 8px;">Quantity Discrepancies</h3>
        <table style="width: 100%; border-collapse: collapse;">
            <thead>
                <tr style="background: #1B365D; color: white;">
                    <th style="padding: 10px; text-align: left;">Description</th>
                    <th style="padding: 10px; text-align: right;">Required</th>
                    <th style="padding: 10px; text-align: right;">Proposed</th>
                    <th style="padding: 10px; text-align: right;">Variance</th>
                </tr>
            </thead>
            <tbody>
'''
            for item in discrepancies[:10]:
                variance = item.get('variance_pct', 0)
                var_color = '#dc3545' if variance < 0 else '#ff9800'
                html += f'''                <tr style="border-bottom: 1px solid #ddd;">
                    <td style="padding: 10px;">{item.get('description', '')[:40]}</td>
                    <td style="padding: 10px; text-align: right;">{item.get('required_qty', 0)}</td>
                    <td style="padding: 10px; text-align: right;">{item.get('proposed_qty', 0)}</td>
                    <td style="padding: 10px; text-align: right; color: {var_color};">{variance:+.1f}%</td>
                </tr>
'''
            html += '''            </tbody>
        </table>
    </div>
'''
        
        # Recommendations
        recommendations = analysis.get('recommendations', [])
        if recommendations:
            html += '''
    <div class="section" style="margin: 20px 0;">
        <h3 style="color: #1B365D; border-bottom: 2px solid #1B365D; padding-bottom: 8px;">Recommendations</h3>
        <ol style="padding-left: 20px;">
'''
            for rec in recommendations[:8]:
                html += f'            <li style="padding: 5px 0;">{rec}</li>\n'
            html += '''        </ol>
    </div>
'''
        
        html += '''
    <div class="footer" style="text-align: center; color: #999; font-size: 12px; margin-top: 30px; padding-top: 20px; border-top: 1px solid #ddd;">
        Generated by Bid Proposal Agent - Abonmarche
    </div>
</div>
'''
        
        return html
